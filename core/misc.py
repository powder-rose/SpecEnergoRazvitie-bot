"""Общая бизнес-логика Telegram-бота и будущего HTTP API."""

from __future__ import annotations

import shutil
import base64
from copy import deepcopy
import logging
import os
import re
import threading
import time
from datetime import datetime
from io import BytesIO
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile
from uuid import uuid4

import pymupdf
import requests
from lxml import etree
from docx import Document
from PIL import Image, ImageOps, UnidentifiedImageError
from docx.opc.part import PartFactory
from docx.parts.document import DocumentPart

DOCM_CONTENT_TYPE = "application/vnd.ms-word.document.macroEnabled.main+xml"
PartFactory.part_type_for[DOCM_CONTENT_TYPE] = DocumentPart


LOGGER = logging.getLogger(__name__)


class ConfigurationError(RuntimeError):
    """Приложение настроено некорректно."""


class DocumentExtractionError(RuntimeError):
    """Из документа не удалось извлечь пригодный текст."""


class AIServiceError(RuntimeError):
    """Ошибка при обращении к Yandex Cloud."""


class MissingCompanyDetailsError(ValueError):
    """В ответе ИИ отсутствуют обязательные реквизиты контрагента."""

    def __init__(self, missing_fields: list[str]) -> None:
        self.missing_fields = tuple(missing_fields)
        super().__init__(
            "Не хватает обязательных реквизитов: "
            + ", ".join(self.missing_fields)
        )


def configure_logging(log_dir: Path, level: str = "INFO") -> None:
    """Настраивает консольный лог и ежедневный ограниченный файл."""
    root_logger = logging.getLogger()
    if getattr(root_logger, "_contract_bot_configured", False):
        return

    log_dir.mkdir(parents=True, exist_ok=True)
    resolved_level = getattr(logging, level.upper(), logging.INFO)
    formatter = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)

    file_handler = RotatingFileHandler(
        log_dir / "contract_bot.log",
        maxBytes=5 * 1024 * 1024,
        backupCount=5,
        encoding="utf-8",
    )
    file_handler.setFormatter(formatter)

    root_logger.setLevel(resolved_level)
    root_logger.addHandler(console_handler)
    root_logger.addHandler(file_handler)
    root_logger._contract_bot_configured = True


class Miscellaneous:
    """Общие операции с файлами, Yandex Cloud и шаблоном договора."""

    DOCM_MAIN_CONTENT_TYPE = (
        "application/vnd.ms-word.document.macroEnabled.main+xml"
    )
    ALLOWED_EXTENSIONS = frozenset({".pdf", ".doc", ".docx", ".docm"})
    MAX_FILE_BYTES = 20 * 1024 * 1024
    MAX_DOCUMENT_CHARS = 100_000
    MAX_SCAN_PAGES = 10
    GPT_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    OCR_URL = "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"
    REQUEST_TIMEOUT = (30, 180)
    OCR_REQUEST_TIMEOUT = (120, 240)
    OCR_MAX_ATTEMPTS = 2
    OCR_MAX_IMAGE_BYTES = 8 * 1024 * 1024
    OCR_MAX_PIXELS = 12_000_000
    OCR_MIN_PIXELS = 160_000
    OCR_MIN_SIDE = 100
    OCR_PDF_DPI = 200
    WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    GENITIVUS = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }
    COMPANY_FIELDS = {
        0: "организационно-правовая форма",
        1: "должность руководителя в родительном падеже",
        2: "должность руководителя",
        3: "ФИО руководителя в родительном падеже",
        4: "ФИО руководителя",
        5: "ИНН",
        6: "юридический адрес",
        7: "наименование банка",
        8: "расчётный счёт",
        9: "корреспондентский счёт",
        10: "БИК",
        11: "КПП",
        12: "ОГРН или ОГРНИП",
        13: "наименование организации",
    }

    EMPTY_COMPANY_VALUES = {
        "",
        "-",
        "—",
        "n/a",
        "none",
        "null",
        "нет",
        "нет данных",
        "не найдено",
        "неизвестно",
        "не указано",
        "отсутствует",
    }

    def __init__(
        self,
        *,
        bot_dir: Path | None = None,
        prompts_dir: Path | None = None,
        trust_env: bool = True,
    ) -> None:
        self.CORE_DIR = Path(__file__).resolve().parent
        self.PROJECT_DIR = self.CORE_DIR.parent
        self.BOT_DIR = Path(bot_dir) if bot_dir else self.PROJECT_DIR / "bot"
        self.PROMPTS_DIR = (
            Path(prompts_dir) if prompts_dir else self.CORE_DIR / "prompts"
        )

        self.USERS_DOCS_DIR = self.BOT_DIR / "usersdocs"
        self.DOCS_DIR = self.BOT_DIR / "docs"
        self.DOC_NUMBERS_DIR = self.CORE_DIR / "docs_numers"
        self.COUNT_NUMBERS_DIR = self.CORE_DIR / "counting_numers"
        for directory in (
            self.USERS_DOCS_DIR,
            self.DOCS_DIR,
            self.DOC_NUMBERS_DIR,
            self.COUNT_NUMBERS_DIR,
        ):
            directory.mkdir(parents=True, exist_ok=True)

        self.YANDEX_FOLDER_ID = self._require_env("YANDEX_FOLDER_ID")
        self.YANDEX_API_KEY = self._require_env("YANDEX_API_KEY")
        self.YANDEXGPT_API_KEY = os.getenv(
            "YANDEXGPT_API_KEY",
            self.YANDEX_API_KEY,
        )

        # Промпты загружаются один раз при создании общего сервиса.
        self.prompt_doc = self._load_prompt("prompt_doc.txt")
        self.prompt_img = self._load_prompt("prompt_pdf.txt")
        self.prompt_mes = self._load_prompt("prompt_mes.txt")

        self._http = requests.Session()
        self._http.trust_env = trust_env
        self._number_lock = threading.Lock()
        LOGGER.info(
            "Общий сервис и промпты инициализированы | trust_env=%s",
            trust_env,
        )

    @staticmethod
    def _require_env(name: str) -> str:
        value = os.getenv(name, "").strip()
        if not value:
            raise ConfigurationError(
                f"Не задана обязательная переменная окружения {name}"
            )
        return value

    def _load_prompt(self, filename: str) -> str:
        path = self.PROMPTS_DIR / filename
        try:
            prompt = path.read_text(encoding="utf-8").strip()
        except OSError as exc:
            raise ConfigurationError(
                f"Не удалось прочитать файл промпта: {path}"
            ) from exc
        if not prompt:
            raise ConfigurationError(f"Файл промпта пуст: {path}")
        return prompt

    def download_bot_doc(self, message: Any, bot: Any) -> Path | None:
        """Загружает допустимый документ Telegram во временную папку."""
        document = getattr(message, "document", None)
        if document is None:
            return None

        extension = Path(document.file_name or "").suffix.lower()
        if extension not in self.ALLOWED_EXTENSIONS:
            LOGGER.warning(
                "Отклонён формат документа | user_id=%s | extension=%s",
                getattr(message.from_user, "id", None),
                extension or "<empty>",
            )
            return None

        file_info = bot.get_file(document.file_id)
        content = bot.download_file(file_info.file_path)
        if len(content) > self.MAX_FILE_BYTES:
            raise DocumentExtractionError(
                f"Файл больше {self.MAX_FILE_BYTES // (1024 * 1024)} МБ"
            )

        target = self.USERS_DOCS_DIR / f"{uuid4().hex}{extension}"
        target.write_bytes(content)
        LOGGER.info(
            "Документ загружен | user_id=%s | extension=%s | bytes=%d",
            getattr(message.from_user, "id", None),
            extension,
            len(content),
        )
        return target

    def delete_local_doc(self, path: str | Path | None) -> None:
        if path is None:
            return
        target = Path(path)
        try:
            target.unlink(missing_ok=True)
            LOGGER.debug("Временный файл удалён | path=%s", target)
        except OSError:
            LOGGER.exception("Не удалось удалить временный файл | path=%s", target)

    @classmethod
    def _is_missing_company_value(cls, value: Any) -> bool:
        if value is None:
            return True
        return str(value).strip().casefold() in cls.EMPTY_COMPANY_VALUES

    def _extract_company_data(self, prompt: str, source_text: str) -> list[str]:
        """Дополняет ответ ИИ реквизитами с явными метками из исходного текста."""

        company = [
            re.sub(r"^\d+[.)]\s*", "", line.strip())
            for line in self._send_text_to_gpt(prompt, source_text).splitlines()
        ]

        # Если ИИ вдруг вернул лишние строки, они не должны сдвигать структуру.
        company = company[:len(self.COMPANY_FIELDS)]

        while len(company) < len(self.COMPANY_FIELDS):
            company.append("")

        labeled_numbers = {
            5: r"(?<![\wА-Яа-я])ИНН\s*[:№-]?\s*(\d{10}|\d{12})(?!\d)",
            8: (
                r"(?:расч[её]тн(?:ый|ого)\s+сч[её]т|р/?с)"
                r"\s*[:№-]?\s*(\d{20})(?!\d)"
            ),
            9: (
                r"(?:корр(?:еспондентский)?\.?\s*сч[её]т|к/?с)"
                r"\s*[:№-]?\s*(\d{20})(?!\d)"
            ),
            10: r"(?<![\wА-Яа-я])БИК\s*[:№-]?\s*(\d{9})(?!\d)",
            11: r"(?<![\wА-Яа-я])КПП\s*[:№-]?\s*(\d{9})(?!\d)",
            12: r"(?<![\wА-Яа-я])ОГРН(?:ИП)?\s*[:№-]?\s*(\d{13}|\d{15})(?!\d)",
        }

        restored: list[str] = []

        for index, pattern in labeled_numbers.items():
            if not self._is_missing_company_value(company[index]):
                continue

            match = re.search(pattern, source_text, flags=re.IGNORECASE)

            if match:
                company[index] = match.group(1)
                restored.append(self.COMPANY_FIELDS[index])

        if restored:
            LOGGER.info(
                "Реквизиты восстановлены из исходного текста | fields=%s",
                ", ".join(restored),
            )

        # -------------------------------------------------
        # Восстанавливаем должность руководителя
        # -------------------------------------------------

        normalized_source = " ".join(
            source_text
            .replace("ё", "е")
            .replace("Ё", "Е")
            .casefold()
            .split()
        )

        # Генеральный директор
        if re.search(
                r"\bгенеральн(?:ый|ого)\s+директор(?:а)?\b",
                normalized_source,
        ):

            current_nom = str(company[2] or "").strip().casefold()
            current_gen = str(company[1] or "").strip().casefold()

            # Именительный падеж
            if (
                    self._is_missing_company_value(company[2])
                    or current_nom == "директор"
            ):
                company[2] = "Генеральный директор"

            # Родительный падеж
            if (
                    self._is_missing_company_value(company[1])
                    or current_gen in {
                "директор",
                "директора",
            }
            ):
                company[1] = "Генерального директора"

        # -------------------------------------------------
        # Нормализуем регистр должности руководителя
        # Выполняется ВСЕГДА, а не только для гендиректора
        # -------------------------------------------------

        for index in (1, 2):
            value = str(company[index] or "").strip()

            if value and not self._is_missing_company_value(value):
                company[index] = value[0].upper() + value[1:]

        return company



    def sent_message_to_ai(self, message: str) -> list[str]:
        text = (message or "").strip()
        if not text:
            raise ValueError("Сообщение для обработки пусто")
        return self._extract_company_data(self.prompt_mes, text)

    def sent_doc_to_ai(self, path: str | Path) -> list[str]:
        """
        Извлекает текст документа.

        Большие документы разбиваются на части,
        после чего найденные реквизиты объединяются.
        """
        document_path = Path(path)
        started_at = time.monotonic()

        text = self.extract_document_text(document_path)

        LOGGER.info(
            "Текст документа извлечён | extension=%s | chars=%d | duration=%.2fs",
            document_path.suffix.lower(),
            len(text),
            time.monotonic() - started_at,
        )

        chunks = self._split_large_text(text, chunk_size=70_000)

        if len(chunks) == 1:
            return self._extract_company_data(
                self.prompt_doc,
                chunks[0],
            )

        LOGGER.info(
            "Большой документ разбит на части | chars=%d | chunks=%d",
            len(text),
            len(chunks),
        )

        merged = [""] * len(self.COMPANY_FIELDS)

        for chunk_index, chunk in enumerate(chunks, start=1):
            LOGGER.info(
                "Обработка части документа | chunk=%d/%d | chars=%d",
                chunk_index,
                len(chunks),
                len(chunk),
            )

            extracted = self._extract_company_data(
                self.prompt_doc,
                chunk,
            )

            for index in range(len(self.COMPANY_FIELDS)):
                if index >= len(extracted):
                    continue

                candidate = str(extracted[index] or "").strip()

                if self._is_missing_company_value(candidate):
                    continue

                if self._is_missing_company_value(merged[index]):
                    merged[index] = candidate

        return merged

    def extract_document_text(self, path: str | Path) -> str:
        """
        Извлекает текст из DOCX, DOC или PDF.

        Ограничение в 100 000 символов здесь специально отсутствует:
        большой текст далее разбивается на части в sent_doc_to_ai().
        """
        document_path = Path(path)

        self._validate_document_path(document_path)

        extension = document_path.suffix.lower()

        extractors = {
            ".docx": self._extract_docx_text,
            ".docm": self._extract_docx_text,
            ".pdf": self._extract_pdf_text,
            ".doc": self._extract_doc_text,
        }

        extractor = extractors.get(extension)

        if extractor is None:
            raise DocumentExtractionError(
                f"Формат {extension or '<без расширения>'} не поддерживается"
            )

        try:
            text = extractor(document_path)
        except DocumentExtractionError:
            raise
        except Exception as exc:
            LOGGER.exception(
                "Внутренняя ошибка извлечения текста | extension=%s",
                extension,
            )
            raise DocumentExtractionError(
                f"Не удалось извлечь текст из файла {extension}"
            ) from exc

        text = self._normalize_extracted_text(text)

        if not text:
            if extension == ".pdf":
                raise DocumentExtractionError(
                    "В PDF нет текстового слоя. "
                    "Отправьте его как PDF-скан."
                )

            raise DocumentExtractionError(
                "Документ не содержит распознаваемого текста"
            )

        return text

    @staticmethod
    def _split_large_text(
            text: str,
            chunk_size: int = 70_000,
    ) -> list[str]:
        """
        Разбивает большой текст на части,
        по возможности не разрывая строки.
        """
        text = (text or "").strip()

        if not text:
            return []

        if len(text) <= chunk_size:
            return [text]

        chunks: list[str] = []
        current: list[str] = []
        current_size = 0

        for line in text.splitlines():
            line = line.strip()

            if not line:
                continue

            line_size = len(line) + 1

            # Если одна строка сама длиннее допустимой части.
            if line_size > chunk_size:
                if current:
                    chunks.append("\n".join(current))
                    current = []
                    current_size = 0

                for start in range(0, len(line), chunk_size):
                    chunks.append(
                        line[start:start + chunk_size]
                    )

                continue

            if current and current_size + line_size > chunk_size:
                chunks.append("\n".join(current))
                current = []
                current_size = 0

            current.append(line)
            current_size += line_size

        if current:
            chunks.append("\n".join(current))

        return chunks

    def _validate_document_path(self, path: Path) -> None:
        if not path.is_file():
            raise DocumentExtractionError(f"Файл не найден: {path}")
        if path.suffix.lower() not in self.ALLOWED_EXTENSIONS:
            raise DocumentExtractionError(
                f"Формат {path.suffix or '<без расширения>'} не поддерживается"
            )
        size = path.stat().st_size
        if size == 0:
            raise DocumentExtractionError("Документ пуст")
        if size > self.MAX_FILE_BYTES:
            raise DocumentExtractionError(
                f"Файл больше {self.MAX_FILE_BYTES // (1024 * 1024)} МБ"
            )

    @staticmethod
    def _normalize_extracted_text(text: str) -> str:
        return "\n".join(
            line.strip()
            for line in text.replace("\x00", "").splitlines()
            if line.strip()
        )

    @staticmethod
    def _extract_docx_text(path: Path) -> str:
        path = Path(path)
        temp_path: Path | None = None
        read_path = path

        if path.suffix.lower() == ".docm":
            # python-docx проверяет Content-Type внутри самого архива, поэтому
            # простого переименования расширения недостаточно — нужно подменить
            # тип документа прямо в [Content_Types].xml.
            temp_path = path.with_name(f"{path.stem}_as_docx.docx")
            docm_type = (
                "application/vnd.ms-word.document.macroEnabled.main+xml"
            )
            docx_type = (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document.main+xml"
            )

            with ZipFile(path, "r") as source_zip:
                with ZipFile(temp_path, "w", ZIP_DEFLATED) as target_zip:
                    for item in source_zip.infolist():
                        data = source_zip.read(item.filename)
                        if item.filename == "[Content_Types].xml":
                            data = data.replace(
                                docm_type.encode("utf-8"),
                                docx_type.encode("utf-8"),
                            )
                        target_zip.writestr(item, data)

            read_path = temp_path

        try:
            document = Document(read_path)
            parts: list[str] = []

            parts.extend(paragraph.text for paragraph in document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))

            for section in document.sections:
                parts.extend(p.text for p in section.header.paragraphs)
                parts.extend(p.text for p in section.footer.paragraphs)

            return "\n".join(parts)
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)

    @staticmethod
    def _extract_pdf_text(path: Path) -> str:
        with pymupdf.open(path) as document:
            if document.needs_pass:
                raise DocumentExtractionError("PDF защищён паролем")
            return "\n".join(page.get_text("text") for page in document)

    @staticmethod
    def _extract_doc_text(path: Path) -> str:
        """
        Извлекает текст из устаревшего бинарного DOC.

        Пакету textract на Linux обычно требуется системная утилита antiword.
        Если окружение сервера её не поддерживает, DOC следует предварительно
        конвертировать в DOCX через LibreOffice.
        """
        try:
            import textract
        except ImportError as exc:
            raise DocumentExtractionError(
                "Для формата .doc установите textract и системную утилиту antiword "
                "либо конвертируйте файл в .docx"
            ) from exc

        try:
            content = textract.process(str(path))
        except Exception as exc:
            raise DocumentExtractionError(
                "Не удалось извлечь текст из .doc; попробуйте сохранить его как .docx"
            ) from exc
        return content.decode("utf-8", errors="replace")

    def normalize_jpeg(self, source):
        source = Path(source)
        target = source.with_name(f"{source.stem}_normalized.jpeg")

        try:
            with Image.open(source) as image:
                image = ImageOps.exif_transpose(image)

                if image.mode in ("RGBA", "LA"):
                    background = Image.new("RGB", image.size, "white")
                    alpha = image.getchannel("A")
                    background.paste(image.convert("RGB"), mask=alpha)
                    image = background
                else:
                    image = image.convert("RGB")
                image.save(target, format="JPEG", quality=95)
        except (UnidentifiedImageError, OSError, ValueError) as exc:
            target.unlink(missing_ok=True)
            raise DocumentExtractionError(
                "Изображение повреждено или имеет неподдерживаемый формат. "
                "Отправьте исходный файл повторно."
            ) from exc

        return target

    def _prepare_ocr_jpeg(self, source: Image.Image) -> bytes:
        """Нормализует и сжимает изображение перед передачей в Vision OCR."""
        image = ImageOps.exif_transpose(source)

        if (
            image.width * image.height < self.OCR_MIN_PIXELS
            or min(image.width, image.height) < self.OCR_MIN_SIDE
        ):
            raise DocumentExtractionError(
                "Изображение слишком маленькое для надёжного распознавания. "
                "Отправьте исходный файл или снимок большего размера."
            )

        if image.mode in ("RGBA", "LA"):
            background = Image.new("RGB", image.size, "white")
            background.paste(image.convert("RGB"), mask=image.getchannel("A"))
            image = background
        else:
            image = image.convert("RGB")

        pixels = image.width * image.height
        if pixels > self.OCR_MAX_PIXELS:
            scale = (self.OCR_MAX_PIXELS / pixels) ** 0.5
            size = (
                max(1, round(image.width * scale)),
                max(1, round(image.height * scale)),
            )
            resampling = getattr(Image, "Resampling", Image).LANCZOS
            image = image.resize(size, resampling)

        content = b""
        for quality in (86, 78, 70, 62, 54):
            buffer = BytesIO()
            image.save(buffer, format="JPEG", quality=quality, optimize=True)
            content = buffer.getvalue()
            if len(content) <= self.OCR_MAX_IMAGE_BYTES:
                break

        if len(content) > self.OCR_MAX_IMAGE_BYTES:
            raise DocumentExtractionError(
                "Страница слишком велика для распознавания даже после сжатия"
            )

        LOGGER.info(
            "Изображение подготовлено для OCR | width=%d | height=%d | bytes=%d",
            image.width,
            image.height,
            len(content),
        )
        return content
    
    def sent_image_to_ai(
        self,
        path: str | Path,
        extension: str | None = None,
    ) -> list[str]:
        del extension  # MIME определяется после нормализации изображения.
        image_path = Path(path)
        try:
            with Image.open(image_path) as source:
                content = self._prepare_ocr_jpeg(source)
        except DocumentExtractionError:
            raise
        except (UnidentifiedImageError, OSError, ValueError) as exc:
            raise DocumentExtractionError(
                "Изображение повреждено или имеет неподдерживаемый формат. "
                "Отправьте исходный файл повторно."
            ) from exc

        text = self._recognize_image_bytes(content, "JPEG")
        return self._extract_company_data(self.prompt_img, text)

    def sent_pdf_scan_to_ai(self, path: str | Path) -> list[str]:
        """
        Обрабатывает PDF.

        Сначала пытается использовать текстовый слой PDF.
        Если текстового слоя нет или текста слишком мало —
        распознаёт страницы через OCR.
        """
        pdf_path = Path(path)

        self._validate_document_path(pdf_path)

        if pdf_path.suffix.lower() != ".pdf":
            raise DocumentExtractionError(
                "Для PDF требуется файл .pdf"
            )

        native_pages: list[str] = []

        with pymupdf.open(pdf_path) as document:

            if document.needs_pass:
                raise DocumentExtractionError(
                    "PDF защищён паролем"
                )

            if len(document) == 0:
                raise DocumentExtractionError(
                    "PDF пуст"
                )

            if len(document) > self.MAX_SCAN_PAGES:
                raise DocumentExtractionError(
                    f"В PDF больше {self.MAX_SCAN_PAGES} страниц"
                )

            # ---------------------------------------------
            # 1. Сначала пробуем обычный текст PDF
            # ---------------------------------------------

            for page_number, page in enumerate(
                    document,
                    start=1,
            ):

                page_text = page.get_text("text").strip()

                if page_text:
                    native_pages.append(
                        f"--- Страница {page_number} ---\n"
                        f"{page_text}"
                    )

            native_text = "\n".join(native_pages).strip()

            # Если получили нормальное количество текста,
            # OCR вообще не используем.
            if len(native_text) >= 100:
                LOGGER.info(
                    "PDF содержит текстовый слой | chars=%d",
                    len(native_text),
                )

                return self._extract_company_data(
                    self.prompt_img,
                    native_text,
                )

            # ---------------------------------------------
            # 2. Если текстового слоя нет — используем OCR
            # ---------------------------------------------

            LOGGER.info(
                "Текстовый слой PDF отсутствует или слишком мал. "
                "Запускается OCR."
            )

            recognized_pages: list[str] = []

            for page_number, page in enumerate(
                    document,
                    start=1,
            ):
                LOGGER.info(
                    "Подготовка страницы PDF к OCR | "
                    "page=%d | pages=%d",
                    page_number,
                    len(document),
                )

                pixmap = page.get_pixmap(
                    dpi=self.OCR_PDF_DPI,
                    colorspace=pymupdf.csRGB,
                    alpha=False,
                )

                image = Image.frombytes(
                    "RGB",
                    (pixmap.width, pixmap.height),
                    pixmap.samples,
                )

                recognized = self._recognize_image_bytes(
                    self._prepare_ocr_jpeg(image),
                    "JPEG",
                )

                recognized_pages.append(
                    f"--- Страница {page_number} ---\n"
                    f"{recognized}"
                )

        text = "\n".join(recognized_pages)

        return self._extract_company_data(
            self.prompt_img,
            text,
        )

    def _recognize_image_bytes(self, content: bytes, mime_type: str) -> str:
        if len(content) > self.OCR_MAX_IMAGE_BYTES:
            raise DocumentExtractionError(
                "Изображение превышает внутренний лимит отправки в OCR"
            )
        payload = {
            "mimeType": mime_type,
            "languageCodes": ["ru", "en"],
            "model": "page",
            "content": base64.b64encode(content).decode("ascii"),
        }
        response = self._post_json(
            self.OCR_URL,
            {
                "Content-Type": "application/json",
                "Authorization": f"Api-Key {self.YANDEX_API_KEY}",
                "x-folder-id": self.YANDEX_FOLDER_ID,
                # Документы могут содержать персональные и банковские данные.
                "x-data-logging-enabled": "false",
            },
            payload,
            operation="ocr",
        )
        try:
            text = response["result"]["textAnnotation"]["fullText"].strip()
        except (KeyError, TypeError, AttributeError) as exc:
            raise AIServiceError("OCR вернул ответ неожиданного формата") from exc
        if not text:
            raise DocumentExtractionError(
                "На изображении не удалось распознать текст. Оно может быть слишком "
                "маленьким, размытым или содержать слишком мелкий текст."
            )
        return text

    def _send_text_to_gpt(self, prompt: str, source_text: str) -> str:
        source_text = source_text.strip()
        if not source_text:
            raise DocumentExtractionError("Нет текста для отправки в YandexGPT")

        payload = {
            "modelUri": f"gpt://{self.YANDEX_FOLDER_ID}/yandexgpt/latest",
            "completionOptions": {
                "stream": False,
                "temperature": 0.1,
            },
            "messages": [
                {
                    "role": "user",
                    "text": f"{prompt}\n\n--- ИСХОДНЫЕ ДАННЫЕ ---\n{source_text}",
                }
            ],
        }
        response = self._post_json(
            self.GPT_URL,
            {
                "Content-Type": "application/json",
                "Authorization": f"Api-Key {self.YANDEXGPT_API_KEY}",
                "x-folder-id": self.YANDEX_FOLDER_ID,
            },
            payload,
            operation="yandexgpt",
        )
        try:
            result = response["result"]["alternatives"][0]["message"]["text"]
        except (KeyError, IndexError, TypeError) as exc:
            raise AIServiceError(
                "YandexGPT вернул ответ неожиданного формата"
            ) from exc
        if not isinstance(result, str) or not result.strip():
            raise AIServiceError("YandexGPT вернул пустой результат")
        return result.strip()

    def _post_json(
        self,
        url: str,
        headers: dict[str, str],
        payload: dict[str, Any],
        *,
        operation: str,
    ) -> dict[str, Any]:
        started_at = time.monotonic()
        is_ocr = operation == "ocr"
        attempts = self.OCR_MAX_ATTEMPTS if is_ocr else 1
        timeout = self.OCR_REQUEST_TIMEOUT if is_ocr else self.REQUEST_TIMEOUT

        for attempt in range(1, attempts + 1):
            attempt_started_at = time.monotonic()
            try:
                response = self._http.post(
                    url,
                    headers=headers,
                    json=payload,
                    timeout=timeout,
                )
                response.raise_for_status()
                data = response.json()
            except (requests.Timeout, requests.ConnectionError) as exc:
                if attempt < attempts:
                    LOGGER.warning(
                        "Сбой соединения с внешним API; запрос будет повторён | "
                        "operation=%s | attempt=%d/%d | duration=%.2fs",
                        operation,
                        attempt,
                        attempts,
                        time.monotonic() - attempt_started_at,
                    )
                    time.sleep(3)
                    continue
                LOGGER.error(
                    "Внешний API недоступен | operation=%s | attempts=%d | "
                    "duration=%.2fs",
                    operation,
                    attempts,
                    time.monotonic() - started_at,
                )
                service = "Yandex OCR" if is_ocr else "YandexGPT"
                raise AIServiceError(
                    f"Не удалось связаться с {service}. "
                    "Проверьте интернет/VPN и повторите попытку."
                ) from exc
            except requests.HTTPError as exc:
                status = exc.response.status_code if exc.response is not None else None
                retryable = status == 429 or (status is not None and status >= 500)
                if retryable and attempt < attempts:
                    LOGGER.warning(
                        "Внешний API временно отклонил запрос; повтор | "
                        "operation=%s | status=%s | attempt=%d/%d",
                        operation,
                        status,
                        attempt,
                        attempts,
                    )
                    time.sleep(3)
                    continue
                LOGGER.error(
                    "Ошибка внешнего API | operation=%s | status=%s | duration=%.2fs",
                    operation,
                    status,
                    time.monotonic() - started_at,
                )
                raise AIServiceError(
                    f"Сервис ИИ отклонил запрос (HTTP {status}). "
                    "Проверьте ключ и права сервисного аккаунта."
                ) from exc
            except requests.RequestException as exc:
                LOGGER.error(
                    "Ошибка внешнего API | operation=%s | duration=%.2fs",
                    operation,
                    time.monotonic() - started_at,
                )
                raise AIServiceError("Ошибка при обращении к сервису ИИ") from exc
            except ValueError as exc:
                raise AIServiceError("Сервис ИИ вернул не JSON") from exc

            LOGGER.info(
                "Внешний API ответил | operation=%s | status=%d | attempt=%d/%d | "
                "duration=%.2fs",
                operation,
                response.status_code,
                attempt,
                attempts,
                time.monotonic() - attempt_started_at,
            )
            return data

        raise AIServiceError("Не удалось выполнить запрос к сервису ИИ")

    def refresh_last_day(self) -> None:
        (self.CORE_DIR / "last_day.txt").write_text(
            datetime.now().strftime("%d%m%y"),
            encoding="utf-8",
        )

    def get_bot_doc_num(
        self,
        user_data: dict[int, dict[str, Any]],
        user_id: int,
    ) -> str:
        return self._get_number(
            user_data,
            user_id,
            self.DOC_NUMBERS_DIR,
            "docs_numers",
            "%d%m%y",
        )

    def get_bot_count_num(
        self,
        user_data: dict[int, dict[str, Any]],
        user_id: int,
    ) -> str:
        return self._get_number(
            user_data,
            user_id,
            self.COUNT_NUMBERS_DIR,
            "countings_numers",
            "%d%m",
        )

    def _get_number(
        self,
        user_data: dict[int, dict[str, Any]],
        user_id: int,
        directory: Path,
        suffix: str,
        date_format: str,
    ) -> str:
        """
        Возвращает следующий номер и сразу резервирует его.

        Формат служебного файла: YYYY-MM-DD|последний_выданный_номер.
        Блокировка защищает от одновременной выдачи одного номера двум потокам.
        """
        initials = str(user_data[user_id]["lastname"])
        counter_path = directory / f"{initials}_{suffix}.txt"
        today = datetime.now().date()

        with self._number_lock:
            number = 1
            try:
                stored_date, stored_number = (
                    counter_path.read_text(encoding="utf-8").strip().split("|", 1)
                )
                if stored_date == today.isoformat():
                    number = int(stored_number) + 1
            except (OSError, ValueError):
                # Старый или повреждённый формат безопасно начинает новый день с 1.
                number = 1

            counter_path.write_text(
                f"{today.isoformat()}|{number}",
                encoding="utf-8",
            )

        result = f"{initials}{datetime.now().strftime(date_format)}-{number}"
        LOGGER.info(
            "Номер зарезервирован | kind=%s | user_id=%s | number=%s",
            suffix,
            user_id,
            result,
        )
        return result

    def integer_texted(self, value: int) -> str:
        prompt = (
            f"Запиши число {value} словами на русском языке. "
            "Верни только результат, первое слово — с большой буквы."
        )
        return self._send_text_to_gpt(prompt, str(value))

    @staticmethod
    def is_individual_entrepreneur(organization_type: str | None) -> bool:
        """Определяет ИП исключительно по организационно-правовой форме."""
        normalized = " ".join(str(organization_type or "").casefold().split())
        return (
            normalized == "ип"
            or normalized.startswith("ип ")
            or "индивидуальный предприниматель" in normalized
            or "индивидуального предпринимателя" in normalized
        )

    @staticmethod
    def abbreviate_fio(full_name: str | None) -> str:
        """
        Иванов Иван Иванович -> Иванов И. И.
        Иванов Иван -> Иванов И.
        """
        value = " ".join(str(full_name or "").strip().split())

        if not value:
            return ""

        parts = value.split()

        if len(parts) == 1:
            return parts[0]

        result = parts[0]

        for part in parts[1:3]:
            clean = part.strip(" .")
            if clean:
                result += f" {clean[0].upper()}."

        return result

    @classmethod
    def build_organization_full_name(
            cls,
            organization_type: str | None,
            organization_name: str | None,
            fio: str | None,
    ) -> str:
        """
        Формирует название заказчика так, как оно должно отображаться в договоре.

        ООО + Ромашка
        -> Общество с ограниченной ответственностью «Ромашка»

        ИП + Иванов Иван Иванович
        -> Индивидуальный предприниматель Иванов Иван Иванович
        """
        org_type = " ".join(str(organization_type or "").strip().split())
        name = " ".join(str(organization_name or "").strip().split())
        person_fio = " ".join(str(fio or "").strip().split())

        if cls.is_individual_entrepreneur(org_type):
            lower_name = name.casefold()

            if lower_name.startswith("индивидуальный предприниматель "):
                return name

            if lower_name.startswith("ип "):
                return (
                        "Индивидуальный предприниматель "
                        + name[3:].strip()
                )

            person = name or person_fio
            return f"Индивидуальный предприниматель {person}".strip()

        opf_map = {
            "ооо": "Общество с ограниченной ответственностью",
            "общество с ограниченной ответственностью":
                "Общество с ограниченной ответственностью",

            "ао": "Акционерное общество",
            "акционерное общество":
                "Акционерное общество",

            "пао": "Публичное акционерное общество",
            "публичное акционерное общество":
                "Публичное акционерное общество",

            "оао": "Открытое акционерное общество",
            "открытое акционерное общество":
                "Открытое акционерное общество",

            "зао": "Закрытое акционерное общество",
            "закрытое акционерное общество":
                "Закрытое акционерное общество",
        }

        # Если ИИ уже вернул полное наименование организации.
        name_lower = name.casefold()

        for short_type, full_type in opf_map.items():
            if name_lower == short_type:
                return full_type

            if name_lower.startswith(short_type + " "):
                tail = name[len(short_type):].strip()

                # ООО "Ромашка" -> Общество с ограниченной
                # ответственностью «Ромашка»
                tail = tail.replace('"', "«", 1)

                if tail.count("«") == 1 and "»" not in tail:
                    tail += "»"

                return f"{full_type} {tail}".strip()

        full_type = opf_map.get(org_type.casefold(), org_type)

        if not name:
            return full_type

        # Если поле 14 содержит только название: Ромашка.
        if (
                "«" not in name
                and "»" not in name
                and '"' not in name
                and org_type.casefold() in opf_map
        ):
            name = f"«{name}»"

        return f"{full_type} {name}".strip()

    @classmethod
    def validate_company_data(
            cls,
            company: list[str | None],
            *,
            raise_on_missing: bool = True,
    ) -> list[str]:
        """
        Возвращает список отсутствующих обязательных реквизитов.

        По умолчанию сохраняется прежнее строгое поведение: при нехватке
        реквизитов возбуждается MissingCompanyDetailsError. Интерфейсы,
        которым нужно сформировать черновик даже из неполных данных (MAX),
        передают raise_on_missing=False и получают список пропусков без
        остановки формирования документа.
        """

        def normalized(index: int) -> str:
            if index >= len(company) or company[index] is None:
                return ""
            return str(company[index]).strip()

        is_entrepreneur = cls.is_individual_entrepreneur(normalized(0))

        required_fields = set(cls.COMPANY_FIELDS)
        required_fields.discard(11)  # КПП не считается обязательным реквизитом

        if is_entrepreneur:
            # У ИП нет отдельного "наименования организации" — есть только ФИО
            required_fields.discard(13)

        missing_fields = [
            cls.COMPANY_FIELDS[index]
            for index in sorted(required_fields)
            if normalized(index).casefold() in cls.EMPTY_COMPANY_VALUES
        ]
        if missing_fields:
            LOGGER.warning(
                "ИИ вернул неполные реквизиты | missing=%s | strict=%s",
                ", ".join(missing_fields),
                raise_on_missing,
            )
            if raise_on_missing:
                raise MissingCompanyDetailsError(missing_fields)

        return missing_fields

    @classmethod
    def _set_replacement_cell_value(
        cls,
        container: Any,
        value: str,
        namespaces: dict[str, str],
    ) -> None:
        """
        Записывает значение в колонку «ИНФОРМАЦИЯ ОТ ЗАКАЗЧИКА».

        Контейнер может быть обычной ячейкой w:tc либо ячейкой,
        обёрнутой в Word Content Control (w:sdt). Оформление таблицы,
        выпадающие списки и соседняя колонка АВТОЗАМЕНА не изменяются.
        """
        w = f"{{{cls.WORD_NS}}}"
        local_name = etree.QName(container).localname

        if local_name == "sdt":
            # После заполнения Word не должен показывать текст-заполнитель
            # «Выберите элемент.».
            for placeholder in container.xpath(
                './w:sdtPr/w:showingPlcHdr',
                namespaces=namespaces,
            ):
                placeholder.getparent().remove(placeholder)

            cells = container.xpath(
                './w:sdtContent/w:tc',
                namespaces=namespaces,
            )
            if not cells:
                raise RuntimeError(
                    "Не найдена ячейка внутри элемента управления Word"
                )
            cell = cells[0]
        else:
            cell = container

        paragraphs = cell.xpath('./w:p', namespaces=namespaces)
        if paragraphs:
            paragraph = paragraphs[0]
            # Убираем лишние параграфы, оставляя геометрию самой ячейки.
            for extra in paragraphs[1:]:
                cell.remove(extra)
        else:
            paragraph = etree.SubElement(cell, w + 'p')

        # Сохраняем оформление исходного абзаца/ячейки (в шаблоне
        # это Times New Roman 9 pt и жёлтая подсветка поля ввода).
        ppr = paragraph.find(w + 'pPr')
        run_props = None
        if ppr is not None:
            ppr_rpr = ppr.find(w + 'rPr')
            if ppr_rpr is not None:
                run_props = deepcopy(ppr_rpr)

        existing_run = paragraph.find(w + 'r')
        if run_props is None and existing_run is not None:
            existing_rpr = existing_run.find(w + 'rPr')
            if existing_rpr is not None:
                run_props = deepcopy(existing_rpr)

        # Удаляем старое содержимое абзаца, но сохраняем pPr.
        for child in list(paragraph):
            if child is not ppr:
                paragraph.remove(child)

        run = etree.SubElement(paragraph, w + 'r')
        if run_props is not None:
            run.append(run_props)

        lines = str(value or '').splitlines() or ['']
        for index, line in enumerate(lines):
            if index:
                etree.SubElement(run, w + 'br')
            text_node = etree.SubElement(run, w + 't')
            if line.startswith(' ') or line.endswith(' '):
                text_node.set(
                    '{http://www.w3.org/XML/1998/namespace}space',
                    'preserve',
                )
            text_node.text = line

    @classmethod
    def _fill_replacement_table(
        cls,
        template_path: Path,
        output_path: Path,
        replacements: dict[str, Any],
        *,
        unwrap_input_controls: bool = False,
    ) -> None:
        """
        Создаёт копию исходной таблицы и заполняет ТОЛЬКО колонку 4.

        Значения autozamena_XXX в колонке 5 остаются буквальным текстом
        и никогда не заменяются. Для шаблона СЭР можно дополнительно
        развернуть content controls колонки 4 в обычные ячейки, чтобы Word
        не обрезал значения по ограничениям старых выпадающих списков.
        """
        namespaces = {'w': cls.WORD_NS}

        with ZipFile(template_path, 'r') as source_archive:
            try:
                document_xml = source_archive.read('word/document.xml')
            except KeyError as exc:
                raise RuntimeError(
                    "В шаблоне отсутствует word/document.xml"
                ) from exc

            root = etree.fromstring(document_xml)
            tables = root.xpath(
                './/w:body/w:tbl',
                namespaces=namespaces,
            )
            if not tables:
                raise RuntimeError("В шаблоне не найдена таблица автозамен")

            # Не предполагаем, что таблица автозамен обязательно первая.
            # Кнопка/элемент управления над таблицей может изменить порядок
            # объектов Word, поэтому ищем таблицу по самим autozamena_XXX.
            expected = set(replacements)

            table = None
            best_matches: set[str] = set()

            for candidate in tables:
                candidate_text = ''.join(
                    candidate.xpath('.//w:t/text()', namespaces=namespaces)
                )
                candidate_matches = set(
                    re.findall(r'autozamena_\d{3}', candidate_text)
                )

                if len(candidate_matches & expected) > len(best_matches & expected):
                    table = candidate
                    best_matches = candidate_matches

                if expected.issubset(candidate_matches):
                    table = candidate
                    best_matches = candidate_matches
                    break

            if table is None or not (best_matches & expected):
                raise RuntimeError(
                    "В шаблоне не найдена таблица с autozamena_001 ... autozamena_020"
                )

            filled: set[str] = set()

            for row in table.xpath('./w:tr', namespaces=namespaces):
                row_text = ''.join(
                    row.xpath('.//w:t/text()', namespaces=namespaces)
                )
                match = re.search(r'autozamena_\d{3}', row_text)
                if not match:
                    continue

                key = match.group(0)
                if key not in replacements:
                    continue

                # Логические ячейки строки. Некоторые поля ввода Word
                # представлены не w:tc напрямую, а w:sdt -> w:tc.
                logical_cells: list[Any] = []
                for child in row:
                    local_name = etree.QName(child).localname
                    if local_name in {'tc', 'sdt'}:
                        logical_cells.append(child)

                if len(logical_cells) < 5:
                    raise RuntimeError(
                        f"Неверная структура строки {key}: "
                        f"ожидалось 5 колонок, найдено {len(logical_cells)}"
                    )

                # Предпоследняя логическая ячейка — строго колонка
                # «ИНФОРМАЦИЯ ОТ ЗАКАЗЧИКА / ДАННЫЕ ДЛЯ ЗАМЕНЫ».
                target_cell = logical_cells[-2]

                # В старом шаблоне СЭР часть ячеек была обёрнута в w:sdt
                # (comboBox). Word повторно применял ограничения этого поля
                # при открытии документа и отрезал последний символ у
                # значений вроде «Генеральный директор», «Устава»,
                # «ежеквартально» и «за три месяца». В рабочей копии СЭР
                # превращаем такую ячейку в обычную w:tc до записи значения.
                if (
                    unwrap_input_controls
                    and etree.QName(target_cell).localname == 'sdt'
                ):
                    actual_cell = target_cell.find(
                        './w:sdtContent/w:tc',
                        namespaces=namespaces,
                    )
                    if actual_cell is None:
                        raise RuntimeError(
                            f"Не удалось развернуть поле ввода для {key}"
                        )
                    plain_cell = deepcopy(actual_cell)
                    target_cell.getparent().replace(target_cell, plain_cell)
                    target_cell = plain_cell

                cls._set_replacement_cell_value(
                    target_cell,
                    str(replacements[key]),
                    namespaces,
                )
                filled.add(key)

            expected = set(replacements)
            missing = sorted(expected - filled)
            if missing:
                raise RuntimeError(
                    "В таблице отсутствуют строки автозамен: "
                    + ", ".join(missing)
                )

            output_path.parent.mkdir(parents=True, exist_ok=True)
            updated_xml = etree.tostring(
                root,
                xml_declaration=True,
                encoding='UTF-8',
                standalone='yes',
            )

            with ZipFile(output_path, 'w', ZIP_DEFLATED) as target_archive:
                for item in source_archive.infolist():
                    data = (
                        updated_xml
                        if item.filename == 'word/document.xml'
                        else source_archive.read(item.filename)
                    )
                    target_archive.writestr(item, data)

    def bot_insert_req(
            self,
            user_data: dict[int, dict[str, Any]],
            user_id: int,
            company: list[str | None],
            numer: str,
            counting: str,
            texted_costs: str,
            texted_total: str,
            way: str | Path | None = None,
            *,
            allow_incomplete: bool = False,
    ) -> Path:
        """
        Формирует НЕ договор, а заполненную таблицу автозамен.

        Исходный template2.docm остаётся визуально неизменным:
        заполняется только колонка «ИНФОРМАЦИЯ ОТ ЗАКАЗЧИКА».
        Текст autozamena_001 ... autozamena_020 не меняется.
        """
        self.validate_company_data(
            company,
            raise_on_missing=not allow_incomplete,
        )

        def field(index: int) -> str:
            value = company[index] if index < len(company) else None
            return str(value).strip() if value is not None else ''

        is_entrepreneur = self.is_individual_entrepreneur(field(0))
        organization_full_name = self.build_organization_full_name(
            field(0),
            field(13),
            field(4),
        )
        fio_short = self.abbreviate_fio(field(4))

        ustav = (
            'листа записи ЕГРИП'
            if is_entrepreneur
            else 'Устава'
        )

        company_req = '\n'.join(
            value
            for value in (
                f'Юридический адрес: {field(6)}' if field(6) else '',
                f'ОГРН: {field(12)}' if field(12) else '',
                f'ИНН: {field(5)}' if field(5) else '',
                (
                    f'КПП: {field(11)}'
                    if not is_entrepreneur and field(11)
                    else ''
                ),
            )
            if value
        )

        banco = '\n'.join(
            value
            for value in (
                f'Банк: {field(7)}' if field(7) else '',
                f'Расчетный счет: {field(8)}' if field(8) else '',
                f'Корр. счет: {field(9)}' if field(9) else '',
                f'БИК: {field(10)}' if field(10) else '',
            )
            if value
        )

        now = datetime.now()
        date_start = (
            f'{now.day} '
            f'{self.GENITIVUS[now.month]} '
            f'{now.year} года'
        )

        replacements = {
            'autozamena_001': numer,
            'autozamena_002': date_start,
            'autozamena_003': user_data[user_id]['ending'],
            'autozamena_004': organization_full_name,
            'autozamena_005': field(1),
            'autozamena_006': field(2),
            'autozamena_007': field(3),
            'autozamena_008': fio_short,
            'autozamena_009': ustav,
            'autozamena_010': field(5),
            'autozamena_011': field(6),
            'autozamena_012': numer,
            'autozamena_013': date_start,
            'autozamena_014': texted_costs,
            'autozamena_015': str(user_data[user_id]['complects']),
            'autozamena_016': str(user_data[user_id]['count_print']),
            'autozamena_017': str(user_data[user_id]['count_sending']),
            'autozamena_018': texted_total,
            'autozamena_019': company_req,
            'autozamena_020': banco,
        }

        template_path = self.CORE_DIR / 'template2.docm'
        if not template_path.is_file():
            raise FileNotFoundError(
                f'Не найден шаблон таблицы: {template_path}'
            )

        source_id = Path(way).stem if way else uuid4().hex
        output_path = (
                self.DOCS_DIR
                / f'ТАБЛИЦА_АВТОЗАМЕНЫ_{source_id}.docm'
        )

        self._fill_replacement_table(
            template_path,
            output_path,
            replacements,
        )

        LOGGER.info(
            'Таблица автозамен сформирована | '
            'user_id=%s | organization=%s | path=%s',
            user_id,
            organization_full_name,
            output_path,
        )
        return output_path
    def bot_insert_req_passport_security(
            self,
            user_data: dict[int, dict[str, Any]],
            user_id: int,
            company: list[str | None],
            numer: str,
            counting: str,
            texted_costs: str,
            way: str | Path | None = None,
            *,
            allow_incomplete: bool = False,
    ) -> Path:
        """Формирует договор ООО СПЕЦКОНС на паспорт безопасности."""
        self.validate_company_data(
            company,
            raise_on_missing=not allow_incomplete,
        )

        def field(index: int) -> str:
            value = company[index] if index < len(company) else None
            return str(value).strip() if value is not None else ''

        is_entrepreneur = self.is_individual_entrepreneur(field(0))
        organization_full_name = self.build_organization_full_name(
            field(0),
            field(13),
            field(4),
        )
        fio_short = self.abbreviate_fio(field(4))
        ustav = 'листа записи ЕГРИП' if is_entrepreneur else 'Устава'

        company_req = '\n'.join(
            value
            for value in (
                f'Юридический адрес: {field(6)}' if field(6) else '',
                f'ОГРН: {field(12)}' if field(12) else '',
                f'ИНН: {field(5)}' if field(5) else '',
                (
                    f'КПП: {field(11)}'
                    if not is_entrepreneur and field(11)
                    else ''
                ),
            )
            if value
        )

        banco = '\n'.join(
            value
            for value in (
                f'Банк: {field(7)}' if field(7) else '',
                f'Расчетный счет: {field(8)}' if field(8) else '',
                f'Корр. счет: {field(9)}' if field(9) else '',
                f'БИК: {field(10)}' if field(10) else '',
            )
            if value
        )

        now = datetime.now()
        date_start = (
            f'{now.day} '
            f'{self.GENITIVUS[now.month]} '
            f'{now.year} года'
        )

        replacements = {
            'autozamena_001': numer,
            'autozamena_002': date_start,
            'autozamena_003': user_data[user_id]['ending'],
            'autozamena_004': organization_full_name,
            'autozamena_005': field(1),
            'autozamena_006': field(2),
            'autozamena_007': field(3),
            'autozamena_008': fio_short,
            'autozamena_009': ustav,
            'autozamena_010': field(5),
            'autozamena_011': field(6),
            'autozamena_012': counting,
            'autozamena_013': str(user_data[user_id]['cost']),
            'autozamena_014': texted_costs,
            'autozamena_019': company_req,
            'autozamena_020': banco,
        }

        template_path = (
            self.CORE_DIR / 'ООО СПЕЦКОНС_ПАСПОРТ БЕЗОПАСНОСТИ.docm'
        )
        if not template_path.is_file():
            raise FileNotFoundError(
                f'Не найден шаблон паспорта безопасности: {template_path}'
            )

        source_id = Path(way).stem if way else uuid4().hex
        output_path = (
            self.DOCS_DIR
            / f'ДОГОВОР_ПАСПОРТ_БЕЗОПАСНОСТИ_{source_id}.docm'
        )

        self._fill_replacement_table(
            template_path,
            output_path,
            replacements,
            unwrap_input_controls=True,
        )

        LOGGER.info(
            'Договор на паспорт безопасности сформирован | '
            'user_id=%s | organization=%s | path=%s',
            user_id,
            organization_full_name,
            output_path,
        )
        return output_path

    @classmethod
    def _validate_ser_template_placeholders(cls, template_path: Path) -> None:
        """Проверяет, что шаблон СЭР использует единую схему autozamena_001..027."""
        namespaces = {'w': cls.WORD_NS}
        with ZipFile(template_path, 'r') as archive:
            try:
                root = etree.fromstring(archive.read('word/document.xml'))
            except KeyError as exc:
                raise RuntimeError(
                    "В шаблоне СЭР отсутствует word/document.xml"
                ) from exc

        document_text = ''.join(
            root.xpath('.//w:t/text()', namespaces=namespaces)
        )
        forbidden = [
            token
            for token in ('autozamena_F013', 'electron_pochta')
            if token in document_text
        ]
        if forbidden:
            raise RuntimeError(
                "В шаблоне СЭР остались устаревшие плейсхолдеры: "
                + ', '.join(forbidden)
            )

        expected = {f'autozamena_{index:03d}' for index in range(1, 28)}
        present = set(re.findall(r'autozamena_\d{3}', document_text))
        missing = sorted(expected - present)
        if missing:
            raise RuntimeError(
                "В шаблоне СЭР отсутствуют плейсхолдеры: "
                + ', '.join(missing)
            )

        # 026/027 должны быть не только в служебной таблице, но и в двух
        # блоках реквизитов Заказчика (основной договор и соглашение ЭДО).
        for key in ('autozamena_026', 'autozamena_027'):
            if document_text.count(key) < 3:
                raise RuntimeError(
                    f"Плейсхолдер {key} не подключён ко всем блокам "
                    "реквизитов Заказчика"
                )

    def bot_insert_req_ser(
            self,
            user_data: dict[int, dict[str, Any]],
            user_id: int,
            company: list[str | None],
            numer: str,
            texted_total: str,
            way: str | Path | None = None,
            *,
            allow_incomplete: bool = False,
    ) -> Path:
        """
        Формирует таблицу автозамен для договора ООО СПЕЦЭНЕРГОРАЗВИТИЕ.

        Использует отдельный шаблон 'ООО СПЕЦЭНЕРГОРАЗВИТИЕ.docm' и
        отдельный набор полей user_data['ser_fields'].
        """
        self.validate_company_data(
            company,
            raise_on_missing=not allow_incomplete,
        )

        def field(index: int) -> str:
            value = company[index] if index < len(company) else None
            return str(value).strip() if value is not None else ''

        is_entrepreneur = self.is_individual_entrepreneur(field(0))
        organization_full_name = self.build_organization_full_name(
            field(0),
            field(13),
            field(4),
        )
        fio_short = self.abbreviate_fio(field(4))

        ustav = (
            'листа записи ЕГРИП'
            if is_entrepreneur
            else 'Устава'
        )

        company_req = '\n'.join(
            value
            for value in (
                f'Юридический адрес: {field(6)}' if field(6) else '',
                f'ОГРН {field(12)}' if field(12) else '',
                f'ИНН {field(5)}' if field(5) else '',
                (
                    f'КПП {field(11)}'
                    if not is_entrepreneur and field(11)
                    else ''
                ),
            )
            if value
        )

        banco = '\n'.join(
            value
            for value in (
                'Полное наименование банка',
                field(7),
                f'РС {field(8)}' if field(8) else '',
                f'КС {field(9)}' if field(9) else '',
                f'БИК {field(10)}' if field(10) else '',
            )
            if value
        )

        now = datetime.now()
        date_start = (
            f'{now.day} '
            f'{self.GENITIVUS[now.month]} '
            f'{now.year} года'
        )

        ser_fields = user_data[user_id].get('ser_fields', {})

        months_count_raw = ser_fields.get('months_count', '')
        cost_month_raw = ser_fields.get('cost_month', '')

        try:
            cost_total_value = int(months_count_raw) * int(cost_month_raw)
            cost_total = str(cost_total_value)
        except (ValueError, TypeError):
            cost_total = ''

        replacements = {
            'autozamena_001': numer,
            'autozamena_002': date_start,
            'autozamena_003': user_data[user_id]['ending'],
            'autozamena_004': organization_full_name,
            'autozamena_005': field(1),
            'autozamena_006': field(2),
            'autozamena_007': field(3),
            'autozamena_008': fio_short,
            'autozamena_009': ustav,
            'autozamena_010': field(5),
            'autozamena_011': field(6),
            'autozamena_012': self.get_bot_count_num(user_data, user_id),
            'autozamena_013': date_start,
            'autozamena_014': months_count_raw,
            'autozamena_015': cost_month_raw,
            'autozamena_016': cost_total,
            'autozamena_017': texted_total,
            'autozamena_018': ser_fields.get('advance', ''),
            'autozamena_019': ser_fields.get('advance_period', ''),
            'autozamena_020': ser_fields.get('object_address', ''),
            'autozamena_021': ser_fields.get('object_name', ''),
            'autozamena_022': ser_fields.get('service_period', ''),
            'autozamena_023': ser_fields.get('visits_frequency', ''),
            'autozamena_024': ser_fields.get('termination_period', ''),
            'autozamena_025': ser_fields.get('email', ''),
            'autozamena_026': company_req,
            'autozamena_027': banco,
        }

        template_path = self.CORE_DIR / 'ООО СПЕЦЭНЕРГОРАЗВИТИЕ.docm'
        if not template_path.is_file():
            raise FileNotFoundError(
                f'Не найден шаблон таблицы: {template_path}'
            )

        self._validate_ser_template_placeholders(template_path)

        source_id = Path(way).stem if way else uuid4().hex
        output_path = (
                self.DOCS_DIR
                / f'ДОГОВОР_СЕР_{source_id}.docm'
        )

        self._fill_replacement_table(
            template_path,
            output_path,
            replacements,
            unwrap_input_controls=True,
        )

        LOGGER.info(
            'Таблица автозамен СЕР сформирована | '
            'user_id=%s | organization=%s | path=%s',
            user_id,
            organization_full_name,
            output_path,
        )
        return output_path

